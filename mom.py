from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from pypandoc import convert_file

import os

def replace_placeholders(template_path, pdf_output_path, replacements):
    """
    Replace placeholders in a .docx template with provided values and set default font size and bold as default.

    :param template_path: Path to the .docx template.
    :param pdf_output_path: Path to save the edited .pdf.
    :param replacements: Dictionary of placeholders and their corresponding values.
                         Keys are placeholders (e.g., '{dept}', '{date}', etc.).
                         Values are strings or lists (e.g., for {abs}, a list of strings).
    """
    temp_docx_path = "temp_output.docx"  # Temporary path for .docx file

    doc = Document(template_path)

    # Set default font size and bold for the document
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(14)
    font.name = "Calibri"
    font.bold = True

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Split the text around the placeholder
                before_text, _, after_text = paragraph.text.partition(placeholder)
                paragraph.text = before_text.strip()  # Keep only text before the placeholder

                if isinstance(value, list):
                    if placeholder == "{_}":  # Special handling for {_} to add normal bullet points
                        for item in value:
                            bullet_paragraph = paragraph.insert_paragraph_before(style="Normal")
                            bullet_run = bullet_paragraph.add_run(f"\u2022 {item}")  # Add bullet character \u2022
                            bullet_run.font.size = Pt(14)
                            bullet_run.font.name = "Calibri"
                            bullet_run.bold = False  # Ensure {_} text is not bold
                    else:  # Handle list of strings for other placeholders
                        for i, item in enumerate(value):
                            if i == 0:
                                # Add the first item to the same paragraph
                                bold_run = paragraph.add_run(item)
                                bold_run.bold = True
                                bold_run.font.size = Pt(14)
                                bold_run.font.name = "Calibri"
                            else:
                                # Add subsequent items in new paragraphs
                                new_paragraph = paragraph.insert_paragraph_before(style="Normal")
                                new_run = new_paragraph.add_run(item)
                                new_run.bold = True
                                new_run.font.size = Pt(14)
                                new_run.font.name = "Calibri"
                else:  # Handle single string replacements
                    if placeholder == "{start_time}" or placeholder == "{end_time}":
                        # Handle start_time and end_time on the same line
                        paragraph.text += f" {replacements['{start_time}']} - {replacements['{end_time}']}"
                        break
                    bold_run = paragraph.add_run(value)
                    bold_run.bold = True
                    bold_run.font.size = Pt(14)
                    bold_run.font.name = "Calibri"

                # Add any text that was after the placeholder
                if after_text.strip():
                    after_run = paragraph.add_run(after_text.strip())
                    after_run.font.size = Pt(14)
                    after_run.font.name = "Calibri"

                break  # Stop processing this placeholder for the current paragraph

    # Save the modified document as a temporary .docx file
    doc.save(temp_docx_path)

    # Convert the .docx file to .pdf
    convert_file(temp_docx_path, to='pdf', outputfile=pdf_output_path)


    # Remove the temporary .docx file
    os.remove(temp_docx_path)

    print(f"Edited .pdf saved to {pdf_output_path}")


#EXAMPLE USAGE

# if __name__ == "__main__":
    # # Define paths
    # template_path = "template.docx"  # Path to your .docx template
    # pdf_output_path = "output.pdf"  # Path to save the modified .pdf

    # # Placeholder replacements
    # replacements = {
    #     "{dept}": "Design Team",
    #     "{date}": "28/12/2024",  # Must follow DD/MM/YYYY format
    #     "{start_time}": "2:00 PM",
    #     "{end_time}": "4:00 PM",
    #     "{agenda}": "Discuss the upcoming project deadlines.",
    #     "{abs}": ["Anurag Prasoon", "Tarun", "Akshat"],
    #     "{_}": ["Submit your report by 5 PM.", "Attend the team meeting at 3 PM.", "Review and finalize the project proposal."]
    # }

    # # Replace placeholders and generate the output PDF
    # replace_placeholders(template_path, pdf_output_path, replacements)
